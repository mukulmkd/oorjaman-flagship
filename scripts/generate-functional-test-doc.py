#!/usr/bin/env python3
"""
Generate the OorjaMan Complete Functional Testing Guide (Word) in project-docs/.

Written for non-technical testers: descriptive flows, not bare test-case IDs.

Run: `.venv-docgen/bin/python scripts/generate-functional-test-doc.py`
     or `npm run docs:functional-test` (requires python-docx)
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "project-docs" / "OorjaMan-Functional-Test-Spec.docx"

ADMIN_PORTAL_URL = "https://oorjaman-admin.vercel.app"
VENDOR_PORTAL_URL = "https://oorjaman-vendor.vercel.app"
SUPPORT_PORTAL_URL = "https://oorjaman-support.vercel.app"
CUSTOMER_APK_LABEL = "OorjaMan Customer (UAT).apk"
TECHNICIAN_APK_LABEL = "OorjaMan Technician (UAT).apk"
TEST_OTP = "123456"

SEED_USERS: list[tuple[str, str, str, str]] = [
    ("Priya Sharma", "Admin", "900000000101", "Admin portal — primary tester"),
    ("Vikram Mehta", "Admin", "900000000102", "Admin portal — second tester"),
    ("Ananya Reddy", "Support", "900000000111", "Support desk — primary"),
    ("Karthik Nair", "Support", "900000000112", "Support desk — second"),
    ("Gamusa Green Energy", "Vendor", "900000000201", "Approved partner · Gamusa"),
    ("Bharat Sun Systems", "Vendor", "900000000202", "Approved partner · Bharat Sun"),
    ("Amit Das", "Technician", "900000000301", "Technician · Gamusa"),
    ("Sanjay Pillai", "Technician", "900000000302", "Technician · Bharat Sun"),
    ("Ravi Iyer", "Technician", "900000000303", "Technician · Gamusa"),
    ("Deepak Menon", "Technician", "900000000304", "Technician · Gamusa"),
    ("Suresh Babu", "Technician", "900000000305", "Technician · Bharat Sun"),
    ("Manoj Krishnan", "Technician", "900000000306", "Technician · Bharat Sun"),
    ("Raju Mahalingam", "Customer", "900000000401", "Customer · Bengaluru home ready to book"),
    ("Rajesh Kumar", "Customer", "900000000402", "Customer — second account"),
    ("Teammate Customer 1", "Customer", "900000000403", "Parallel testing"),
    ("Teammate Customer 2", "Customer", "900000000404", "Parallel testing"),
    ("Teammate Customer 3", "Customer", "900000000405", "Parallel testing"),
    ("Teammate Customer 4", "Customer", "900000000406", "Parallel testing"),
]


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------


def shade_cell(cell, fill: str = "E8F0FE") -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        shade_cell(table.rows[0].cells[i])
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    doc.add_paragraph()


def add_callout(doc: Document, title: str, lines: list[str], fill: str = "FFF8E1") -> None:
    """Highlighted box for 'What you should see' / tips."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x5D, 0x40, 0x37)
    for line in lines:
        bp = doc.add_paragraph(line, style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.25)


def add_flow_heading(doc: Document, title: str, level: int = 2) -> None:
    doc.add_heading(title, level=level)


def add_narrative(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def add_checkpoints(doc: Document, items: list[str]) -> None:
    doc.add_paragraph("Check that:")
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)


def add_scenario(
    doc: Document,
    title: str,
    who: str,
    overview: str,
    steps: list[str],
    checkpoints: list[str],
) -> None:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(f"Who tests this: {who}")
    add_narrative(doc, overview)
    doc.add_paragraph("What to do:")
    add_steps(doc, steps)
    add_checkpoints(doc, checkpoints)
    doc.add_paragraph()


def add_checkbox_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    return doc


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------


def add_cover_and_intro(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("OorjaMan\nComplete Functional Testing Guide")
    run.bold = True
    run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"For QA testers and business reviewers · {date.today().strftime('%d %B %Y')}")

    doc.add_paragraph()
    add_narrative(
        doc,
        "This guide walks you through every major feature in OorjaMan — the customer mobile app, "
        "technician mobile app, admin portal, vendor (partner) portal, and support desk. "
        "It is written in plain language. You do not need to be a developer to use it."
    )
    add_narrative(
        doc,
        "Read each section in order the first time, then use the end-to-end stories in Part 7 "
        "as play-scripts for release testing. Tick the checkboxes as you go. "
        "Report anything that crashes, looks broken, or does not match what this guide says should happen."
    )

    doc.add_heading("How to use this guide", level=1)
    for item in [
        "Each flow explains what the user sees, what happens behind the scenes in simple terms, "
        "and what you should verify.",
        "Use the test phone numbers in Part 0. The verification code is always "
        f"{TEST_OTP} — no real SMS is sent in the test environment.",
        "Many flows involve more than one app. Coordinate with teammates playing customer, "
        "vendor, technician, admin, and support roles.",
        "When a flow says “Pass if”, treat every bullet as required unless your project lead says otherwise.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_part_0_getting_started(doc: Document) -> None:
    doc.add_heading("Part 0 — Before you begin", level=1)

    doc.add_heading("What you need", level=2)
    for item in [
        "An Android phone for the Customer app and Technician app.",
        "Chrome or Safari on a computer or tablet for the three web portals.",
        f"The two APK files from your project contact: {CUSTOMER_APK_LABEL} and {TECHNICIAN_APK_LABEL}.",
        f"The test verification code: {TEST_OTP}.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Where to open each app", level=2)
    add_table(
        doc,
        ["Application", "Who uses it", "How to open it"],
        [
            ["Customer app", "Homeowners", f"Install {CUSTOMER_APK_LABEL} on Android → open OorjaMan Customer"],
            ["Technician app", "Field technicians", f"Install {TECHNICIAN_APK_LABEL} on Android → open OorjaMan Technician"],
            ["Admin portal", "OorjaMan operations staff", f"Browser → {ADMIN_PORTAL_URL}"],
            ["Vendor portal", "Solar cleaning partners", f"Browser → {VENDOR_PORTAL_URL}"],
            ["Support desk", "Customer & technician support", f"Browser → {SUPPORT_PORTAL_URL}"],
        ],
    )

    doc.add_heading("How to sign in (same on every app)", level=2)
    add_steps(
        doc,
        [
            "Choose India (+91) if the screen asks.",
            "Enter the 10-digit phone number from the table below (example: 900000000401).",
            "Tap Send code or Continue.",
            f"Enter the verification code: {TEST_OTP}.",
            "Tap Verify or Sign in.",
        ],
    )
    add_callout(
        doc,
        "Note:",
        [
            "A test-mode banner on the login screen is normal — it reminds you that no real SMS is sent.",
            "If login fails for everyone, tell your project contact; the test environment may need a refresh.",
        ],
    )

    doc.add_heading("Test accounts", level=2)
    add_table(
        doc,
        ["Name", "Role", "Phone", "Best used for"],
        [[n, r, p, notes] for n, r, p, notes in SEED_USERS],
    )
    add_narrative(
        doc,
        "Raju (401) already completed onboarding with a Bengaluru home address and 5 kW solar — "
        "use him for most customer booking tests. Gamusa technicians (301, 303, 304) work for vendor 201; "
        "Bharat Sun technicians (302, 305, 306) work for vendor 202."
    )


def add_part_1_glossary(doc: Document) -> None:
    doc.add_heading("Part 1 — Words and statuses you will see", level=1)
    add_narrative(
        doc,
        "Bookings and contracts move through stages. The labels below appear on customer phones, "
        "partner portals, and admin screens. Use this table when a status does not match what you expect."
    )

    doc.add_heading("Visit (booking) statuses — customer view", level=2)
    add_table(
        doc,
        ["Status on screen", "What it means", "What usually happens next"],
        [
            ["Awaiting payment", "Visit created but not paid yet", "Customer completes test payment → visit is confirmed"],
            ["Awaiting partner / Partner confirming", "Paid; OorjaMan or partner has not accepted yet", "Partner accepts within 1 hour and assigns a technician"],
            ["Partner acknowledged", "Partner saw the request", "Partner assigns technician"],
            ["Technician assigned", "A named technician is on the job", "Technician marks “On the way” → customer can track"],
            ["Visit in progress", "Technician started work on site", "Technician finishes with photos and codes → Completed"],
            ["Completed", "Cleaning finished", "Customer can rate, view receipt, see photos in Activity"],
            ["Cancelled", "Visit will not happen", "Refund or credit rules apply depending on who cancelled and when"],
        ],
    )

    doc.add_heading("AMC (Annual Maintenance Contract) statuses", level=2)
    add_table(
        doc,
        ["Status", "What the customer sees", "What it means"],
        [
            ["No AMC", "Plan options on AMC tab", "Customer can subscribe to a plan"],
            ["Payment pending", "Complete payment message", "Subscription started but payment not finished"],
            ["Awaiting setup", "“OorjaMan will assign your partner”", "Paid; admin assigns partner and funds visit wallet"],
            ["Active", "Plan name, renewal date, visit slots", "Customer can book included visits"],
            ["Visits used up", "Upgrade or pay for extra visit", "All included visits consumed this contract period"],
            ["Expiring / Renewal due", "Renewal banner on Home or AMC tab", "Customer should renew before end date"],
            ["Expired / Cancelled", "Subscribe again prompt", "No included visits until new contract"],
        ],
    )

    doc.add_heading("Partner (vendor) approval statuses", level=2)
    add_table(
        doc,
        ["Status", "What the partner sees"],
        [
            ["Registration in progress", "Continue signup wizard"],
            ["Awaiting approval", "Submitted; waiting for OorjaMan admin"],
            ["Approved", "Full dashboard access — Operations, Team, Finance, etc."],
            ["Rejected", "Application declined — contact OorjaMan"],
            ["Suspended", "Account paused — cannot accept new visits"],
        ],
    )

    doc.add_heading("Technician onboarding statuses", level=2)
    add_table(
        doc,
        ["Status", "What the technician sees"],
        [
            ["Not invited", "“Ask your employer to add your phone” screen"],
            ["Complete profile", "5-step onboarding wizard"],
            ["Waiting for employer approval", "Static waiting screen with Refresh"],
            ["Approved", "Home and Jobs tabs — can receive assignments"],
            ["Rejected", "Message to contact employer or support"],
        ],
    )

    doc.add_heading("Important time rules (test these carefully)", level=2)
    for item in [
        "Partners must accept or decline an incoming visit within 1 hour of it being assigned.",
        "Customers can cancel free of charge until 1 hour before the scheduled visit window; "
        "after that a late-cancellation fee may apply.",
        "Job Start Code and Happy Code are 6-digit codes shared between customer and technician "
        "to prove the right people met at the right time.",
        "OorjaMan Credits (₹1 per credit) are issued when a partner cancels close to visit time; "
        "credits apply to future one-time visits only.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_part_2_customer_app(doc: Document) -> None:
    doc.add_heading("Part 2 — Customer app (homeowner mobile app)", level=1)
    add_narrative(
        doc,
        "The customer app is for homeowners who book solar panel cleaning. "
        "It has five bottom tabs: Home, Bookings, AMC, Activity, and Profile."
    )

    # --- First launch ---
    doc.add_heading("2.1 First launch — welcome, permissions, and login", level=2)
    add_narrative(
        doc,
        "When someone opens the app for the first time (or after uninstalling), they see a short "
        "introduction before signing in."
    )
    add_scenario(
        doc,
        "Welcome slides",
        "Any tester with a fresh install",
        "Four slides explain OorjaMan: easy booking, solar cleaning, and why use the service.",
        [
            "Install and open the Customer app.",
            "Swipe through all welcome slides.",
            "Tap Get started on the last slide.",
        ],
        [
            "Slides are readable and images are not cut off.",
            "You reach the location permission screen or login.",
        ],
    )
    add_scenario(
        doc,
        "Location permission",
        "Any tester",
        "Location helps show nearby partners and lets technicians find the site. It can be skipped initially.",
        [
            "On the permissions screen, tap Allow while using the app (or Skip / Not now if offered).",
            "If you skipped, note that booking may ask for location again later.",
        ],
        [
            "App continues to login without looping back to welcome.",
            "Choice is remembered for this install.",
        ],
    )
    add_scenario(
        doc,
        "Login with test account",
        "Use Raju · 900000000401",
        "Existing customers sign in with phone + verification code.",
        [
            f"Enter phone 900000000401.",
            f"Tap Send code, then enter {TEST_OTP}.",
            "Tap Verify.",
        ],
        [
            "Home tab opens with bottom navigation (Home, Bookings, AMC, Activity, Profile).",
            "App does not crash after login.",
            "Profile shows Raju’s name and phone.",
        ],
    )
    add_scenario(
        doc,
        "Wrong app — vendor or technician phone",
        "Use vendor 900000000201 or technician 900000000301",
        "Each phone number is tied to one role. The customer app must reject other roles.",
        [
            "Try to log into the customer app with a vendor or technician test number.",
        ],
        [
            "A clear message says this account is not for the customer app.",
            "You cannot reach the Home screen.",
        ],
    )

    # --- Registration ---
    doc.add_heading("2.2 New customer registration (5 steps)", level=2)
    add_narrative(
        doc,
        "A brand-new phone number (not in the test table) is sent to the registration wizard after first login. "
        "The main app tabs stay blocked until all five steps are finished."
    )
    add_steps(
        doc,
        [
            "Step 1 — About you: Enter display name, contact email, and optional alternate phone.",
            "Step 2 — Address: Enter site label (e.g. Home), street, city, district, state, pincode.",
            "Step 3 — Location: Tap to capture GPS location on the roof/site (grant permission if asked).",
            "Step 4 — Solar & site: Enter capacity (kW), panel count, residential/commercial, roof type, "
            "access, water availability, hazards, brands, last cleaning date, special instructions.",
            "Step 5 — Safety & terms: Tick all three consent boxes (accurate info, terms/privacy, contact permission) "
            "and submit.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Each step validates required fields — you cannot skip empty mandatory items.",
            "Back button keeps entered data where possible.",
            "After submit, you land on the Home tab.",
            "If you sign out and back in, you go straight to Home (not registration again).",
        ],
    )
    add_callout(
        doc,
        "Mandatory address gate:",
        [
            "Some accounts must pick or confirm a default service address every session until one is saved.",
            "A full-screen picker blocks the tabs until you select or add an address.",
        ],
    )

    # --- Home tab ---
    doc.add_heading("2.3 Home tab", level=2)
    add_narrative(
        doc,
        "Home is the starting point. It shows a greeting, your service location, quick actions to book a visit, "
        "AMC status cards (if applicable), and renewal prompts when a contract is ending."
    )
    add_checkpoints(
        doc,
        [
            "Service location picker in the header switches between saved addresses.",
            "Book a visit button is visible when online.",
            "AMC active or renew cards appear when relevant.",
            "Help / support button in the header opens support chat.",
            "With airplane mode on, an offline message appears and booking is blocked politely.",
        ],
    )

    # --- One-time booking ---
    doc.add_heading("2.4 Booking a one-time (paid) cleaning visit", level=2)
    add_narrative(
        doc,
        "A one-time visit is a single paid cleaning. The booking wizard has four steps: Partner → Schedule → "
        "Confirm → Payment."
    )

    doc.add_heading("Step-by-step: one-time booking wizard", level=3)
    add_steps(
        doc,
        [
            "From Home or Bookings, tap Book a visit.",
            "If you have an active AMC, the app may first ask: use an included AMC visit or book a paid one-time visit. "
            "Choose one-time / paid visit.",
            "Step 1 — Partner: Pick a preferred partner (e.g. Gamusa Green Energy) or Any OorjaMan partner. "
            "Confirm which service address this visit is for.",
            "Step 2 — Schedule: Pick a date (up to 90 days ahead) and an available time slot.",
            "Step 3 — Confirm: Review address, recipient (yourself or someone else), price with GST breakdown, "
            "and add special notes if needed.",
            "Step 4 — Payment: Complete payment in test mode (no real money). You may apply OorjaMan Credits "
            "if you have a balance.",
        ],
    )

    doc.add_heading("What happens after you book (full lifecycle)", level=3)
    add_narrative(
        doc,
        "After payment, the visit appears under Bookings. Here is the full journey from booking to completion:"
    )
    add_table(
        doc,
        ["Stage", "What the customer sees", "What others do", "What to verify"],
        [
            ["1. Just booked", "Status: Awaiting partner", "OorjaMan routes to preferred or nearest partner", "Booking listed under Bookings; price shown"],
            ["2. Partner accepts", "Technician assigned (name may appear)", "Vendor accepts in portal + assigns technician", "Push notification; status updates"],
            ["3. On the way", "Technician on the way + Track button", "Technician taps En route in their app", "Live map on Track screen refreshes"],
            ["4. On site", "Job Start Code visible on booking detail", "Technician enters your Job Start Code", "Code matches what technician app asks for"],
            ["5. In progress", "Visit in progress", "Technician takes before/after photos", "Status updates without manual refresh"],
            ["6. Complete", "Completed + Happy Code + rate prompt", "Technician enters Happy Code and finishes", "Receipt/share option; Activity shows photos"],
        ],
    )
    add_callout(
        doc,
        "Job Start Code and Happy Code:",
        [
            "Job Start Code — shown on your booking when the technician is due; give this to the technician to start.",
            "Happy Code — shown during/after the visit; share with the technician to confirm completion.",
            "You can regenerate a code if needed (short cooldown applies).",
            "Use the share button to send codes via WhatsApp or SMS.",
        ],
    )

    # --- AMC ---
    doc.add_heading("2.5 AMC — subscribe, manage, and book included visits", level=2)
    add_narrative(
        doc,
        "AMC (Annual Maintenance Contract) covers multiple scheduled cleanings per year for one service address. "
        "Each address can have its own AMC."
    )

    doc.add_heading("Subscribing to AMC (first time)", level=3)
    add_steps(
        doc,
        [
            "Open the AMC tab.",
            "Select the service address (if you have more than one).",
            "Choose a plan tier based on your solar capacity (kW).",
            "Review price and complete test payment.",
            "After payment you see: OorjaMan will assign your dedicated partner.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "AMC tab shows plan name and status (awaiting setup or active after admin assigns partner).",
            "Home may show an AMC active card.",
            "Admin AMC Contracts page shows the new contract.",
        ],
    )

    doc.add_heading("After AMC is active — booking an included visit", level=3)
    add_narrative(
        doc,
        "Included visits do not require payment. The wizard skips the payment step."
    )
    add_steps(
        doc,
        [
            "On AMC tab, tap Book next visit (or Book a visit from Home when AMC slot is available).",
            "Choose partner and schedule as usual.",
            "Confirm — no payment screen.",
            "Visit appears under Bookings marked as AMC-covered.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Visit slot on AMC tab moves from Not scheduled → Scheduled → Completed after the job.",
            "No duplicate charge for the included visit.",
        ],
    )

    doc.add_heading("AMC updates — upgrade, renewal, and capacity changes", level=3)
    add_table(
        doc,
        ["Action", "Where", "What happens"],
        [
            ["Upgrade plan", "AMC tab → Upgrade", "Higher tier; more visits; may require top-up payment"],
            ["Renew contract", "Renewal banner on Home/AMC when expiring", "Extend end date; payment in test mode"],
            ["Change solar kW in Profile", "Profile → solar details", "App may warn that AMC tier/pricing could change"],
            ["All visits used", "AMC tab", "Prompt to upgrade or book a paid one-time visit"],
            ["Partner not yet assigned", "AMC tab after subscribe", "Cannot book included visit until admin assigns partner"],
        ],
    )
    add_scenario(
        doc,
        "AMC renewal flow",
        "Customer 402 or account near expiry",
        "When a contract is ending, renewal prompts appear on Home and AMC tab.",
        [
            "Open the renewal prompt or AMC tab.",
            "Review plan and price.",
            "Complete renewal payment in test mode.",
        ],
        [
            "Contract end date extends.",
            "Activity tab shows a renewal event.",
            "Admin Subscription Renewals page can show reminder history.",
        ],
    )

    # --- Cancel reschedule ---
    doc.add_heading("2.6 Reschedule and cancel a visit", level=2)

    doc.add_heading("Reschedule", level=3)
    add_steps(
        doc,
        [
            "Open Bookings → tap the visit → Reschedule.",
            "Pick a new date (within the allowed window, typically 14 days) and time slot.",
            "Confirm the change.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "New date and time show on booking detail.",
            "Vendor and technician apps show the updated schedule.",
            "If too close to visit time, reschedule may be blocked with a clear message.",
        ],
    )

    doc.add_heading("Cancel — free vs late cancellation", level=3)
    add_narrative(
        doc,
        "Cancellation is allowed while the visit is still upcoming (before work starts). "
        "Whether a fee applies depends on how close you are to the scheduled time."
    )
    add_steps(
        doc,
        [
            "Open booking detail for an upcoming visit.",
            "Tap Cancel.",
            "Select a reason from the list.",
            "If a late-cancellation fee applies, the app shows the amount and asks you to accept before confirming.",
            "Confirm cancellation.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Status changes to Cancelled on Bookings list and detail.",
            "Free cancel: no fee message when more than 1 hour before the visit window.",
            "Late cancel: fee amount shown and must be accepted.",
            "If partner cancelled near visit time, you may receive OorjaMan Credits (check Credits wallet).",
        ],
    )

    # --- After completion ---
    doc.add_heading("2.7 After a visit completes", level=2)
    add_checkpoints(
        doc,
        [
            "Booking detail shows Completed status.",
            "Rate the visit (stars + optional comment) — prompt should not repeat after rating.",
            "Share service receipt / tax summary if offered.",
            "Activity tab shows the visit with before/after site photos.",
            "Push notification received for completion.",
        ],
    )

    # --- Profile ---
    doc.add_heading("2.8 Profile — view and update your information", level=2)
    add_narrative(
        doc,
        "Everything entered during registration can be updated later from the Profile tab."
    )
    add_table(
        doc,
        ["Section", "What you can change", "What stays fixed"],
        [
            ["Personal", "Display name, contact email, alternate phone", "Login phone number"],
            ["Addresses", "Add/edit/remove service addresses; set default site", "—"],
            ["Solar & site", "kW, panels, roof, access, brands, instructions", "May trigger AMC tier warning"],
            ["Preferred partners", "Choose favourite partners per address", "From partners you have used"],
            ["Site photos", "Upload and view stamped site photos", "—"],
            ["OorjaMan Credits", "View balance and transaction history", "Credits applied automatically at checkout"],
            ["AMC summary", "View active plans per address", "Subscribe/upgrade from AMC tab"],
            ["Sign out", "Ends session", "—"],
        ],
    )
    add_scenario(
        doc,
        "Add a second address and book from it",
        "Raju 401",
        "Customers with multiple sites can switch addresses when booking.",
        [
            "Profile → Addresses → Add new address (e.g. Office).",
            "Set as default if desired.",
            "Book a visit and confirm the new address is pre-selected.",
        ],
        [
            "Both addresses appear in the address picker.",
            "Booking is created for the correct address.",
        ],
    )

    # --- Other tabs ---
    doc.add_heading("2.9 Bookings, Activity, and support tabs", level=2)
    add_narrative(doc, "Bookings tab lists upcoming and past visits with status chips. Tap any row for detail.")
    add_narrative(
        doc,
        "Activity tab is a timeline per address: status changes, reschedules, AMC events, and ratings."
    )
    add_scenario(
        doc,
        "Support chat",
        "Any logged-in customer",
        "Customers can get help without leaving the app.",
        [
            "Tap the help/support button from Home, Bookings, AMC, Activity, or Profile.",
            "Choose a category: Booking related, AMC related, or Any other query.",
            "Pick a sub-topic (e.g. schedule change, renewal, payment).",
            "Optionally link a booking or AMC contract.",
            "Send a message and wait for support reply.",
        ],
        [
            "Chat thread opens and messages send successfully.",
            "After support resolves the case, you may see a short satisfaction survey.",
        ],
    )

    doc.add_heading("2.10 Notifications and real-time updates", level=2)
    add_checkpoints(
        doc,
        [
            "Notification when partner accepts and technician is assigned.",
            "Notification when technician is on the way.",
            "Notification when visit is completed.",
            "Tapping a notification opens the correct booking.",
            "Booking detail updates live when technician changes status (no need to pull-to-refresh constantly).",
            "AMC renewal or offer banners appear when relevant.",
        ],
    )

    doc.add_heading("2.11 Preferred partners", level=2)
    add_narrative(
        doc,
        "After you have completed at least one visit with a partner, you can mark them as preferred "
        "for a specific service address. Future bookings can pre-select that partner."
    )
    add_steps(
        doc,
        [
            "Profile → Preferred partners (or open from booking wizard Step 1).",
            "Select a service address.",
            "Add Gamusa Green Energy or Bharat Sun Systems from your history or the directory.",
            "Save preference.",
            "Start a new booking — preferred partner should appear as default choice.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Preference is saved per address (not globally).",
            "Choosing Any OorjaMan partner lets OorjaMan auto-route if preferred partner is unavailable.",
        ],
    )

    doc.add_heading("2.12 Bookings tab — upcoming vs past", level=2)
    add_checkpoints(
        doc,
        [
            "Upcoming shows active visits (awaiting partner through in progress).",
            "Past shows completed and cancelled visits.",
            "Status chip on each row matches booking detail.",
            "Empty state shows helpful message when no bookings.",
            "“OorjaMan is assigning your partner” message when visit not yet routed.",
        ],
    )


def add_part_3_technician_app(doc: Document) -> None:
    doc.add_heading("Part 3 — Technician app (field partner mobile app)", level=1)
    add_narrative(
        doc,
        "Technicians work for a solar cleaning partner (vendor). They see assigned jobs, navigate to sites, "
        "and complete a structured on-site checklist. Tabs: Home, Jobs, Feedback, Activity, Profile."
    )

    doc.add_heading("3.1 First launch and login", level=2)
    add_scenario(
        doc,
        "Welcome and login",
        "Amit Das · 900000000301",
        "Same pattern as customer app: welcome slides → permissions → login.",
        [
            "Open Technician app (fresh install optional for welcome test).",
            "Complete welcome slides and permissions.",
            f"Login with 900000000301 and OTP {TEST_OTP}.",
        ],
        [
            "Home tab opens for a fully approved technician.",
            "Employer shows Gamusa Green Energy.",
        ],
    )

    doc.add_heading("3.2 Gates — screens that block access", level=2)
    add_table(
        doc,
        ["Screen", "When you see it", "How to clear it"],
        [
            ["Wrong role", "Customer/vendor phone used", "Use a technician test number"],
            ["Vendor not onboarded", "Phone not invited by any partner", "Vendor must invite from Team tab"],
            ["Waiting for employer approval", "Profile submitted, vendor not approved yet", "Vendor approves in Team tab → Refresh"],
            ["Complete profile", "Invited but onboarding not finished", "Finish 5-step wizard and submit"],
        ],
    )

    doc.add_heading("3.3 New technician onboarding (5 steps)", level=2)
    add_narrative(
        doc,
        "When a vendor invites a new phone number, that person completes onboarding before seeing jobs."
    )
    add_steps(
        doc,
        [
            "Vendor invites phone from Vendor portal → Team tab (see Part 4).",
            "Technician installs app and logs in with invited phone.",
            "Step 1 — Employer & you: Employer is pre-filled; enter personal details, address, DOB, emergency contact.",
            "Step 2 — Identity: PAN, Aadhaar last 4, upload ID documents and photo.",
            "Step 3 — Skills: Select skills, work city, service radius, solar experience.",
            "Step 4 — Safety: Complete safety training acknowledgements.",
            "Step 5 — Bank: Account details and bank proof upload → Submit.",
            "Wait on “Waiting for employer approval” screen.",
            "Vendor approves in Team tab → technician taps Refresh → Home opens.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Can save draft and continue later on each step.",
            "After approval, OorjaMan partner ID (employee code) is visible on profile.",
            "Jobs tab becomes active.",
        ],
    )

    doc.add_heading("3.4 Jobs — from assignment to completion", level=2)
    add_narrative(
        doc,
        "Jobs appear after the vendor assigns a visit. The list is grouped into Today, Upcoming, and Completed."
    )

    doc.add_heading("Job detail — before leaving for site", level=3)
    add_checkpoints(
        doc,
        [
            "Customer address and navigation link (open in Google Maps).",
            "Scheduled date and time slot.",
            "Site photos from customer profile if available.",
            "Job Start Code section when visit is accepted.",
        ],
    )

    doc.add_heading("En route and live tracking", level=3)
    add_steps(
        doc,
        [
            "Open assigned job.",
            "Tap En route (location permission required).",
            "Travel to customer site.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Customer app shows Technician on the way and Track map updates.",
            "Technician location continues updating while en route (background).",
        ],
    )

    doc.add_heading("On-site execution wizard (8 steps)", level=3)
    add_narrative(
        doc,
        "At the site, the technician runs a guided checklist. This is the core field workflow."
    )
    add_steps(
        doc,
        [
            "Tap Start visit to open the execution wizard.",
            "Step 1 — Job Start Code: Enter the code from the customer’s booking detail.",
            "Step 2 — Safety confirmations: Tick all required safety items.",
            "Step 3 — Start selfie: Take a photo of yourself on site.",
            "Step 4 — Start timer: Tap to officially start the visit (status becomes In progress).",
            "Step 5 — Before photos: Capture at least one before-cleaning photo.",
            "Step 6 — After photos: Capture after-cleaning photos.",
            "Step 7 — Issues & notes: Record any issues (optional).",
            "Step 8 — Happy Code & finish: Enter the customer’s Happy Code and submit.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Wrong Job Start Code shows an error — visit does not start.",
            "Wrong Happy Code shows an error — visit stays in progress.",
            "After success, job moves to Completed.",
            "Customer receives completion notification.",
            "If visit was already in progress, wizard resumes at the correct step.",
        ],
    )

    doc.add_heading("3.5 Profile, availability, and documents", level=2)
    add_checkpoints(
        doc,
        [
            "Profile shows name, employer, OorjaMan ID, ratings summary.",
            "Available for assignments toggle saves on/off state.",
            "Documents screen shows uploaded IDs; can upload more.",
            "Sign out returns to login and stops location tracking.",
        ],
    )

    doc.add_heading("3.6 Technician support and notifications", level=2)
    add_checkpoints(
        doc,
        [
            "Support button available from Jobs, Profile, and during execution wizard.",
            "New job assignment appears in Jobs list without manual refresh.",
            "Feedback tab shows customer ratings after completed jobs.",
            "Activity tab lists job events.",
            "Bharat Sun technician (302) only sees Bharat Sun jobs — not Gamusa jobs.",
        ],
    )


def add_part_4_vendor_portal(doc: Document) -> None:
    doc.add_heading("Part 4 — Vendor portal (partner web app)", level=1)
    add_narrative(
        doc,
        f"Solar cleaning companies use the vendor portal at {VENDOR_PORTAL_URL}. "
        "Tabs after approval: Overview, Operations, Insights, Finance, Team, Bookings (history), Service coverage."
    )

    doc.add_heading("4.1 Login and account states", level=2)
    add_scenario(
        doc,
        "Approved partner login",
        "Gamusa · 900000000201",
        "Approved partners reach the full dashboard.",
        [
            f"Open {VENDOR_PORTAL_URL}/login.",
            f"Sign in with 900000000201 and OTP {TEST_OTP}.",
            "Open Overview tab.",
        ],
        [
            "Dashboard loads with summary numbers.",
            "All main tabs are accessible.",
        ],
    )
    add_table(
        doc,
        ["Account state", "What the partner sees", "Full dashboard?"],
        [
            ["Not registered", "Link to start registration", "No"],
            ["Registration in progress", "Continue signup wizard", "No"],
            ["Awaiting approval", "Waiting message + Refresh", "No"],
            ["Approved", "Overview and all tabs", "Yes"],
            ["Rejected / Suspended", "Status message and next steps", "No"],
        ],
    )

    doc.add_heading("4.2 New partner registration (8 sections)", level=2)
    add_narrative(
        doc,
        "A new solar company applies through the signup wizard. Progress can be saved and resumed."
    )
    add_steps(
        doc,
        [
            f"Open {VENDOR_PORTAL_URL}/signup.",
            "Section 1 — Partner login: Email for future portal login.",
            "Section 2 — Company: Business name, trade name, GSTIN, PAN, company type.",
            "Section 3 — Contact person: Name, role, phone, email.",
            "Section 4 — Address & service areas: Registered address and regions served.",
            "Section 5 — Experience & workforce: Company summary, years, team size.",
            "Section 6 — Equipment & safety: Equipment list and safety/insurance confirmations.",
            "Section 7 — Bank details: Bank name, IFSC, account number.",
            "Section 8 — Documents: Upload PAN, Aadhaar, GST certificate, bank proof, optional logo.",
            "Submit application.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Validation errors show if required fields are missing.",
            "After submit, portal shows Awaiting approval.",
            "Admin sees application in Vendor approval queue (Part 5).",
        ],
    )

    doc.add_heading("4.3 Operations — accepting visits and assigning technicians", level=2)
    add_narrative(
        doc,
        "This is the heart of partner day-to-day work. Incoming visits must be accepted within 1 hour."
    )

    doc.add_heading("Decline an incoming visit", level=3)
    add_steps(
        doc,
        [
            "Open Operations → Incoming visit.",
            "Tap Decline or Reject.",
            "Enter a reason and confirm.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Visit leaves your incoming queue.",
            "Admin Operations desk may show reassignment or marketplace float.",
            "Customer may see a different partner assigned later.",
        ],
    )

    doc.add_heading("Incoming visit — accept and assign", level=3)
    add_steps(
        doc,
        [
            "Open Operations tab.",
            "Find visit under Incoming (status: awaiting partner).",
            "Tap Accept.",
            "Immediately assign a verified technician (e.g. Amit 301 for Gamusa).",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Visit moves to Active visits.",
            "Technician app shows the job.",
            "Customer app shows Technician assigned.",
        ],
    )

    doc.add_heading("Marketplace visits", level=3)
    add_narrative(
        doc,
        "If the preferred partner does not respond in time, or admin floats the visit, it appears in Marketplace "
        "for any eligible partner to claim."
    )
    add_steps(
        doc,
        [
            "Ensure Available for marketplace claims is turned on (if shown).",
            "Open Marketplace panel in Operations.",
            "Review visit details.",
            "Tap Claim, then Accept and assign technician.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Claimed visit follows same flow as direct assignment.",
            "Customer sees updated partner name.",
        ],
    )

    doc.add_heading("Active visit — monitor and cancel", level=3)
    add_checkpoints(
        doc,
        [
            "Active visits show en route and in progress statuses.",
            "Can view customer site photos.",
            "Partner can cancel an accepted visit with reason (may trigger penalty settlement).",
            "History tab shows completed and cancelled visits with penalty notes where applicable.",
        ],
    )

    doc.add_heading("4.4 Team — invite and approve technicians", level=2)
    add_steps(
        doc,
        [
            "Open Team tab.",
            "Tap Invite technician.",
            "Enter full name, phone (10+ digits), optional email.",
            "Send invite.",
            "Technician completes onboarding on mobile app (Part 3).",
            "Back on Team tab, find pending profile → Approve or Reject.",
            "On approve, technician’s OorjaMan ID appears.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Invite list shows statuses: invited, opened, completed, expired, cancelled.",
            "Can revoke a pending invite.",
            "Approved technicians appear in assignment dropdown on Operations.",
        ],
    )

    doc.add_heading("4.5 Finance, Insights, Coverage, and History", level=2)
    add_checkpoints(
        doc,
        [
            "Finance tab: pending payouts, settled amounts, cancellation penalties; export CSV.",
            "Insights tab: charts for SLA and completion rates.",
            "Service coverage tab: edit pincodes / areas served.",
            "Bookings (history) tab: detailed visit history with photos.",
            "Overview tab KPIs update when new visits are routed (real-time).",
            "Notification bell shows new incoming visit alerts.",
        ],
    )


def add_part_5_admin_portal(doc: Document) -> None:
    doc.add_heading("Part 5 — Admin portal (OorjaMan operations)", level=1)
    add_narrative(
        doc,
        f"OorjaMan staff use the admin portal at {ADMIN_PORTAL_URL} to route visits, "
        "approve partners, manage AMC contracts, finance, and platform settings."
    )

    doc.add_heading("5.1 Login and navigation", level=2)
    add_scenario(
        doc,
        "Admin login",
        "Priya Sharma · 900000000101",
        "Admins land on the Operations desk by default.",
        [
            f"Open {ADMIN_PORTAL_URL}/login.",
            f"Sign in with 900000000101 and OTP {TEST_OTP}.",
            "Confirm Operations desk loads.",
        ],
        [
            "Sidebar shows all sections: Operations, Bookings, Booking routing, Finance, AMC contracts, "
            "Vendor approval, Vendors, Technicians, Partner quality, Trust & safety, Analytics, "
            "Notifications, Feature management, Brand collateral, Renewal reminders, Service pricing.",
            "Support phone (111) cannot access admin — access denied.",
        ],
    )

    doc.add_heading("5.2 Operations desk — daily monitoring", level=2)
    add_narrative(
        doc,
        "The Operations desk is the unified inbox for issues needing attention."
    )
    add_table(
        doc,
        ["Queue type", "What it means", "Typical action"],
        [
            ["Partner response overdue", "Partner did not accept within 1 hour", "Reassign, float to marketplace, or contact partner"],
            ["Marketplace float", "Visit available for any partner to claim", "Monitor until claimed"],
            ["Reassignment needed", "Visit needs a new partner or technician", "Assign from modal"],
            ["AMC awaiting partner", "New AMC paid but no partner assigned", "Assign partner on AMC contracts page"],
            ["OTP lock reset", "Too many wrong code attempts", "Reset for customer or technician"],
            ["Field exception", "Technician flagged an issue on site", "Review and resolve"],
        ],
    )
    add_checkpoints(
        doc,
        [
            "KPI cards show counts and load without endless spinner.",
            "Clicking a queue item opens detail with booking context.",
            "Can assign vendor shortcut from operations modal.",
        ],
    )

    doc.add_heading("5.3 Bookings and booking routing", level=2)
    add_narrative(
        doc,
        "Booking monitoring lists all visits with filters. Booking routing is where unassigned visits get a partner."
    )
    add_steps(
        doc,
        [
            "Open Booking routing.",
            "Find an unassigned customer visit (e.g. Raju 401’s new booking).",
            "Select Gamusa Green Energy (or Bharat Sun).",
            "Save assignment.",
            "Open Bookings monitoring → search by phone 900000000401.",
            "Open booking detail → verify status timeline and OTP codes match customer/technician apps.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Vendor portal shows incoming visit.",
            "Can reassign to a different partner.",
            "Admin can cancel a visit if needed — customer app shows Cancelled.",
            "Admin can reschedule a visit when the feature is available.",
            "OTP lock reset: if customer or technician entered wrong codes too many times, reset from Operations desk.",
        ],
    )

    doc.add_heading("5.4 Vendor approval", level=2)
    add_steps(
        doc,
        [
            "Open Vendor approval queue.",
            "Open a submitted intake application.",
            "Review all 8 sections and uploaded documents.",
            "Approve → partner can log in to full dashboard.",
            "Or Reject with reason → partner sees rejection on portal.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Approved vendors appear in All vendors → approved tab.",
            "Vendor detail page shows organisation profile.",
            "Can suspend a partner (blocks new routing).",
        ],
    )

    doc.add_heading("5.5 AMC contracts and renewals", level=2)
    add_steps(
        doc,
        [
            "Open Finance → AMC contracts.",
            "Filter: pending funding, funded, depleted, cancelled.",
            "For new subscription: assign partner to customer’s AMC.",
            "Confirm wallet funded and visit slots unlocked.",
            "Open Subscription renewals → view expiring/lapsed customers.",
            "Send or verify renewal reminder (email/SMS/WhatsApp in test mode).",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Customer AMC tab shows active status after partner assignment.",
            "Customer can book included visit.",
            "Wallet balance matches visit allowance.",
        ],
    )

    doc.add_heading("5.6 Finance and settlements", level=2)
    add_checkpoints(
        doc,
        [
            "Finance page lists vendor settlement batches.",
            "Line items per visit payout and cancellation penalty.",
            "Statuses: pending review, approved, settled, waived.",
            "Mark settlement processed → vendor Finance tab updates.",
            "Platform fee configuration visible.",
        ],
    )

    doc.add_heading("5.7 Technicians, quality, trust, analytics, and settings", level=2)
    add_checkpoints(
        doc,
        [
            "Technician directory lists all seed technicians with verification badges.",
            "Technician detail shows documents, skills, employer link.",
            "Partner quality dashboard shows vendor scorecards.",
            "Trust & safety page lists incidents/flags.",
            "Analytics charts for bookings and revenue with date range.",
            "Notification templates — view/edit booking and renewal templates.",
            "Feature management — toggle feature flags (test propagation to apps).",
            "Service pricing — AMC plans, one-time rates, capacity slots.",
            "Brand collateral — download letterhead, business cards, etc.",
            "Notification bell in header shows recent ops events.",
        ],
    )


def add_part_6_support_desk(doc: Document) -> None:
    doc.add_heading("Part 6 — Support desk (customer & technician help)", level=1)
    add_narrative(
        doc,
        f"Support agents use {SUPPORT_PORTAL_URL} to answer chats from customers and technicians."
    )

    doc.add_heading("6.1 Login and Insights dashboard", level=2)
    add_scenario(
        doc,
        "Support login",
        "Ananya Reddy · 900000000111",
        "Insights is the home page with live queue metrics.",
        [
            f"Open {SUPPORT_PORTAL_URL}/login.",
            f"Sign in with 900000000111 and OTP {TEST_OTP}.",
            "Review Insights tiles: open chats, queued, unassigned, resolved, CSAT.",
        ],
        [
            "Metrics load without error.",
            "Clicking a tile filters the inbox (e.g. Booking category).",
        ],
    )

    doc.add_heading("6.2 Inbox — handle a conversation", level=2)
    add_narrative(
        doc,
        "The inbox is where agents read and reply to chats. Queue tabs: Queued, Unassigned, Mine, All open, Resolved."
    )
    add_steps(
        doc,
        [
            "Open Inbox.",
            "Filter audience: All, Customer support, or Technician support.",
            "Select a conversation from the list.",
            "Read message history (updates in real time).",
            "Review participant panel: customer/technician profile, bookings, AMC, vendor.",
            "Type a reply and Send.",
            "Use Macros picker for canned responses if helpful.",
            "Add an internal note (not visible to customer) if needed.",
            "Claim or assign the conversation if unassigned.",
            "When done: Resolve → choose Resolved, Duplicate, or Policy limitation.",
            "Customer receives CSAT survey after resolve.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Messages sent from support appear in customer app chat.",
            "Attachments can be uploaded.",
            "Chat dock allows handling multiple conversations.",
            "Sound chime on new message (can be muted).",
            "Resolved conversations leave active queue but stay in Resolved (30d) tab.",
            "Can reopen a resolved conversation if the customer writes again or issue persists.",
        ],
    )

    doc.add_heading("6.3 Search", level=2)
    add_steps(
        doc,
        [
            "Open Search.",
            "Search customers by name, email, or phone (e.g. 900000000401).",
            "Search technicians by name, employee code, or vendor.",
            "Search conversations by subject, category, or booking reference.",
        ],
    )
    add_checkpoints(
        doc,
        [
            "Results open profile with chat history link.",
            "Can jump from search result into inbox thread.",
        ],
    )

    doc.add_heading("6.4 Support categories (what customers pick in the app)", level=2)
    add_table(
        doc,
        ["Category", "Example topics"],
        [
            ["Booking related", "Schedule change, technician issue, payment/refund, wrong address, tracking"],
            ["AMC related", "Plan upgrade, urgent paid cleaning, schedule AMC visit, renewal, visit allowance, pricing"],
            ["Any other query", "Profile/login, addresses, app bug, notifications, general"],
        ],
    )


def add_part_7_e2e_stories(doc: Document) -> None:
    doc.add_heading("Part 7 — End-to-end play-scripts (do once per release)", level=1)
    add_narrative(
        doc,
        "These stories tie all five apps together. Assign roles to different testers and run each script completely."
    )

    stories = [
        (
            "Story A — One-time visit from booking to rating",
            "Customer 401 · Admin 101 · Vendor 201 · Technician 301",
            [
                "Customer Raju books a one-time paid cleaning (Partner → Schedule → Confirm → Pay).",
                "Admin Priya routes the visit to Gamusa on Booking routing.",
                "Vendor Gamusa accepts and assigns Amit in Operations.",
                "Technician Amit taps En route, then completes the 8-step execution wizard with Job Start and Happy codes.",
                "Customer tracks on map, sees completion, rates the visit.",
                "All parties check their History/Activity/Bookings views.",
            ],
            [
                "End status: Completed everywhere.",
                "Photos visible in customer Activity.",
                "Vendor History and admin monitoring match.",
            ],
        ),
        (
            "Story B — AMC subscribe → included visit → complete",
            "Customer 402 · Admin 101 · Vendor 202 · Technician 302",
            [
                "Customer 402 subscribes to AMC on AMC tab and pays.",
                "Admin assigns partner Bharat Sun on AMC contracts and confirms wallet funded.",
                "Customer books included AMC visit (no payment step).",
                "Vendor Bharat Sun accepts and assigns Sanjay 302.",
                "Technician completes visit.",
            ],
            [
                "AMC visit slot shows Completed on AMC tab.",
                "No duplicate one-time charge.",
                "Admin AMC contract shows correct wallet debit.",
            ],
        ),
        (
            "Story C — Reschedule and late cancel",
            "Customer 401 · Vendor 201",
            [
                "Customer books a visit for tomorrow.",
                "Customer reschedules to a different day from booking detail.",
                "Verify vendor and technician see new time.",
                "Book another visit; cancel within free window — no fee.",
                "Book another; cancel inside penalty window — accept fee and confirm.",
            ],
            [
                "Reschedule reflected on all apps.",
                "Free cancel: no fee shown.",
                "Late cancel: fee amount displayed and recorded.",
            ],
        ),
        (
            "Story D — Partner cancel near visit → customer credits",
            "Customer 401 · Vendor 201",
            [
                "Customer has upcoming visit within 1 hour window.",
                "Vendor cancels accepted visit from Operations with reason.",
                "Customer checks OorjaMan Credits wallet.",
            ],
            [
                "Visit shows Cancelled for customer.",
                "Credits issued per policy (if applicable).",
                "Vendor Finance may show cancellation penalty.",
            ],
        ),
        (
            "Story E — New vendor onboarding",
            "New vendor phone · Admin 101",
            [
                f"Complete 8-section signup at {VENDOR_PORTAL_URL}/signup.",
                f"Admin reviews and approves on Vendor approval.",
                "Vendor logs in and reaches Overview.",
                "Admin routes a test visit to new vendor (if service area matches).",
            ],
            [
                "Pending → approved transition works.",
                "Dashboard tabs accessible after approval.",
            ],
        ),
        (
            "Story F — New technician invite and approval",
            "Vendor 201 · New technician phone",
            [
                "Vendor invites new phone from Team tab.",
                "Technician installs app, logs in, completes 5-step onboarding.",
                "Vendor approves on Team tab.",
                "Vendor assigns a real visit to new technician.",
                "Technician completes execution wizard.",
            ],
            [
                "Waiting screen before approval.",
                "Full job access after approval.",
            ],
        ),
        (
            "Story G — Marketplace fallback",
            "Customer 401 · Admin 101 · Vendor 202",
            [
                "Customer books with preferred partner Gamusa.",
                "Admin floats visit to marketplace (or wait for SLA breach in test).",
                "Vendor Bharat Sun claims from Marketplace and assigns technician.",
            ],
            [
                "Customer sees updated partner.",
                "Visit completes normally.",
            ],
        ),
        (
            "Story H — AMC renewal",
            "Customer 401 · Admin 101",
            [
                "Use account with expiring AMC (or seed data).",
                "Customer sees renewal banner; completes renewal payment.",
                "Admin verifies on Subscription renewals page.",
            ],
            [
                "Contract end date extended.",
                "Included visits available for new period.",
            ],
        ),
        (
            "Story I — Profile and AMC upgrade",
            "Customer 401",
            [
                "Profile: update display name and add second address.",
                "AMC tab: upgrade to higher tier if available.",
                "Book visit from new address.",
            ],
            [
                "Profile changes persist after app restart.",
                "Upgrade reflected on AMC tab.",
            ],
        ),
        (
            "Story J — Support escalation",
            "Customer 401 · Support 111 · Admin 101",
            [
                "Customer opens support chat about active booking.",
                "Support claims conversation, replies, links booking context.",
                "If needed, escalate to admin Operations desk.",
                "Admin takes action on booking; support resolves with Resolved tag.",
            ],
            [
                "Customer sees resolution message.",
                "CSAT survey appears.",
                "Case in Resolved tab.",
            ],
        ),
        (
            "Story K — Settlements after completed visits",
            "Admin 101 · Vendor 201",
            [
                "Complete several visits for Gamusa (Story A multiple times or use seed data).",
                "Admin opens Finance settlements for Gamusa.",
                "Review and mark settlement processed.",
                "Vendor checks Finance tab.",
            ],
            [
                "Amounts match completed visit tariffs.",
                "Vendor sees settled status.",
            ],
        ),
    ]

    for title, who, steps, checkpoints in stories:
        add_scenario(doc, title, who, "", steps, checkpoints)


def add_part_8_checklists(doc: Document) -> None:
    doc.add_heading("Part 8 — Master checklists", level=1)

    doc.add_heading("8.1 Customer app — complete checklist", level=2)
    add_checkbox_list(
        doc,
        [
            "First launch welcome and permissions",
            "Login and wrong-role rejection",
            "New customer 5-step registration",
            "Mandatory address gate",
            "Home tab and offline behaviour",
            "One-time booking — full 4-step wizard and payment",
            "Post-booking lifecycle: awaiting partner → assigned → en route → in progress → completed",
            "Track technician map",
            "Job Start Code and Happy Code display and share",
            "Rate completed visit",
            "AMC subscribe and awaiting-setup state",
            "AMC included visit booking (no payment)",
            "AMC upgrade and renewal",
            "Reschedule visit",
            "Free cancel and late cancel with fee",
            "Profile updates: name, addresses, solar details, preferred partners",
            "Site photo gallery",
            "OorjaMan Credits wallet",
            "Bookings, AMC, Activity tabs",
            "Support chat all categories",
            "Push notifications and live status updates",
            "Sign out and session restore",
        ],
    )

    doc.add_heading("8.2 Technician app — complete checklist", level=2)
    add_checkbox_list(
        doc,
        [
            "Welcome, permissions, login",
            "Wrong role and gate screens",
            "5-step onboarding wizard",
            "Pending employer approval and refresh",
            "Jobs list: today, upcoming, completed",
            "Job detail and navigation",
            "En route and customer tracking",
            "8-step execution wizard (all steps)",
            "Invalid code rejection",
            "Profile, availability toggle, documents",
            "Support from job screens",
            "Realtime new assignment",
            "Vendor isolation (Gamusa vs Bharat Sun)",
            "Sign out",
        ],
    )

    doc.add_heading("8.3 Vendor portal — complete checklist", level=2)
    add_checkbox_list(
        doc,
        [
            "Login approved partner",
            "Pending/rejected/suspended states",
            "8-section signup and draft save",
            "Overview KPIs",
            "Operations: accept, assign, active visits",
            "Marketplace claim",
            "Cancel visit with reason",
            "Team: invite, approve, reject, revoke",
            "Finance settlements and export",
            "Insights, Coverage, Bookings history",
            "Notifications",
        ],
    )

    doc.add_heading("8.4 Admin portal — complete checklist", level=2)
    add_checkbox_list(
        doc,
        [
            "Login and sidebar navigation",
            "Operations desk queues and modals",
            "Booking monitoring and search",
            "Booking routing assign and reassign",
            "Vendor approval approve/reject",
            "All vendors and vendor detail",
            "AMC contracts assign partner and wallet",
            "Subscription renewals",
            "Finance settlements",
            "Technician directory",
            "Partner quality, Trust & safety, Analytics",
            "Notification templates, Feature flags",
            "Service pricing and brand collateral",
        ],
    )

    doc.add_heading("8.5 Support desk — complete checklist", level=2)
    add_checkbox_list(
        doc,
        [
            "Login and Insights",
            "Inbox queues and audience filters",
            "Reply, macros, internal notes",
            "Claim, assign, resolve with tags",
            "CSAT after resolve",
            "Search customers, technicians, conversations",
            "Chat dock and realtime updates",
            "Customer chat parity with mobile app",
        ],
    )

    doc.add_heading("8.6 Release sign-off", level=2)
    add_narrative(doc, "Complete before UAT exit:")
    add_checkbox_list(
        doc,
        [
            "All Part 7 end-to-end stories (A through K) executed or waived with ticket",
            "All Part 8 master checklists reviewed",
            "Customer and Technician APKs install on Android",
            "All three web portals open in browser",
            f"OTP {TEST_OTP} works on every app",
            "No crash immediately after login on any app",
            "P1 issues: zero open blockers",
            "QA lead sign-off",
            "Product owner sign-off",
        ],
    )
    add_table(
        doc,
        ["Role", "Name", "Date", "Signature"],
        [["QA Lead", "", "", ""], ["Product Owner", "", "", ""], ["Business reviewer", "", "", ""]],
    )

    doc.add_heading("When something goes wrong", level=2)
    for item in [
        "Note which app, which test phone, and what you tapped just before the problem.",
        "Take a screenshot or screen recording.",
        "If login fails for everyone, contact your project lead — the test environment may need a refresh.",
        "To see first-time welcome slides again, uninstall the mobile app and reinstall the APK.",
        "To reset test bookings, ask your project lead to refresh seed data.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    note = doc.add_paragraph(
        f"Portal URLs: Admin {ADMIN_PORTAL_URL} · Vendor {VENDOR_PORTAL_URL} · Support {SUPPORT_PORTAL_URL}"
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = setup_document()
    add_cover_and_intro(doc)
    add_part_0_getting_started(doc)
    add_part_1_glossary(doc)
    add_part_2_customer_app(doc)
    add_part_3_technician_app(doc)
    add_part_4_vendor_portal(doc)
    add_part_5_admin_portal(doc)
    add_part_6_support_desk(doc)
    add_part_7_e2e_stories(doc)
    add_part_8_checklists(doc)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print("Format: non-technical descriptive flows across all 5 apps")


if __name__ == "__main__":
    main()
