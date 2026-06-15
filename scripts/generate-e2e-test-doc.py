#!/usr/bin/env python3
"""
Generate the non-technical UAT test guide (Word) in the repo root.

Run: python3 scripts/generate-e2e-test-doc.py
Requires: python-docx (use a venv: pip install python-docx)
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "OorjaMan-E2E-Test-Guide.docx"

# --- UAT access (update here when URLs or APK names change) ---
ADMIN_PORTAL_URL = "https://oorjaman-admin.vercel.app"
VENDOR_PORTAL_URL = "https://oorjaman-vendor.vercel.app"
SUPPORT_PORTAL_URL = "https://oorjaman-support.vercel.app"

CUSTOMER_APK_LABEL = "OorjaMan Customer (UAT).apk"
TECHNICIAN_APK_LABEL = "OorjaMan Technician (UAT).apk"

TEST_OTP = "123456"


def shade_cell(cell, fill: str = "E8F0FE") -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        shade_cell(table.rows[0].cells[i])
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    doc.add_paragraph()


def add_test(
    doc: Document,
    number: str,
    title: str,
    steps: list[str],
    pass_criteria: list[str],
) -> None:
    p = doc.add_paragraph()
    p.add_run(f"Test {number}: {title}").bold = True
    doc.add_paragraph("What to do:")
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph("Pass if:")
    for item in pass_criteria:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("OorjaMan\nUAT Testing Guide")
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"For QA testers · {date.today().strftime('%d %B %Y')}")

    doc.add_paragraph(
        "Use this guide to test OorjaMan before release. You will use test phone numbers "
        "and a fixed verification code—no real SMS is sent. Report any screen that crashes, "
        "looks broken, or does not match the “Pass if” notes below."
    )

    # --- Access ---
    doc.add_heading("1. What you need", level=1)
    for item in [
        "An Android phone (for the two mobile apps).",
        "Chrome or Safari on a computer or phone (for the three web portals).",
        f"The test verification code: {TEST_OTP} (use this every time the app asks for an OTP).",
        "The two APK files from your project contact (customer app + technician app).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Where to open each app", level=1)
    add_table(
        doc,
        ["App", "How to open it"],
        [
            ["Customer app (homeowners)", f"Install {CUSTOMER_APK_LABEL} on your Android phone, then open OorjaMan Customer."],
            ["Technician app (field partners)", f"Install {TECHNICIAN_APK_LABEL} on your Android phone, then open OorjaMan Technician."],
            ["Admin portal (operations)", f"Browser → {ADMIN_PORTAL_URL}"],
            ["Vendor portal (solar partners)", f"Browser → {VENDOR_PORTAL_URL}"],
            ["Support desk", f"Browser → {SUPPORT_PORTAL_URL}"],
        ],
    )
    doc.add_paragraph(
        "Installing an APK: transfer the file to your phone, tap it, and allow “Install from unknown sources” if Android asks. "
        "If you already have an older test build, uninstall it first or install over it when prompted."
    )

    doc.add_heading("3. How to sign in (all apps)", level=1)
    doc.add_paragraph("The steps are the same on mobile and web:")
    for step in [
        "Choose India (+91) if asked.",
        "Enter the 10-digit test phone number from the table below (example: 9000000401 — do not type +91 in the box unless the screen asks for full number).",
        "Tap Send code (or Continue).",
        f"Enter OTP: {TEST_OTP}.",
        "Tap Verify or Sign in.",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "A yellow or grey banner on the login screen may remind you that test mode is on. That is expected."
    )

    doc.add_heading("4. Test accounts", level=1)
    doc.add_paragraph("Use these people as examples. Each row is a ready-made account in the test system.")

    add_table(
        doc,
        ["Who", "Role", "Phone to enter", "Best used for"],
        [
            ["Priya Sharma", "Admin", "900000000101", "Admin portal"],
            ["Vikram Mehta", "Admin", "900000000102", "Admin portal (second tester)"],
            ["Ananya Reddy", "Support", "900000000111", "Support desk"],
            ["Karthik Nair", "Support", "900000000112", "Support desk (second tester)"],
            ["Gamusa Green Energy", "Vendor", "900000000201", "Vendor portal"],
            ["Bharat Sun Systems", "Vendor", "900000000202", "Vendor portal (second partner)"],
            ["Amit Das", "Technician", "900000000301", "Technician app · Gamusa jobs"],
            ["Sanjay Pillai", "Technician", "900000000302", "Technician app · Bharat Sun jobs"],
            ["Ravi Iyer", "Technician", "900000000303", "Technician app · Gamusa"],
            ["Deepak Menon", "Technician", "900000000304", "Technician app · Gamusa"],
            ["Suresh Babu", "Technician", "900000000305", "Technician app · Bharat Sun"],
            ["Manoj Krishnan", "Technician", "900000000306", "Technician app · Bharat Sun"],
            ["Raju Mahalingam", "Customer", "900000000401", "Customer app · main booking tests"],
            ["Rajesh Kumar", "Customer", "900000000402", "Customer app · second customer"],
            ["Teammate Customer 1", "Customer", "900000000403", "Parallel testing with a colleague"],
            ["Teammate Customer 2", "Customer", "900000000404", "Parallel testing"],
            ["Teammate Customer 3", "Customer", "900000000405", "Parallel testing"],
            ["Teammate Customer 4", "Customer", "900000000406", "Parallel testing"],
        ],
    )

    doc.add_paragraph(
        "Raju (401) already has a home address in Bengaluru and can book visits straight away. "
        "Gamusa technicians (301, 303, 304) belong to vendor 201; Bharat Sun technicians (302, 305, 306) belong to vendor 202."
    )

    doc.add_heading("5. Testing a brand-new customer (first-time setup)", level=1)
    doc.add_paragraph(
        "Most test customers above are already set up. To see the full new-user journey, either uninstall the customer app "
        "and reinstall it, or ask your project contact for a fresh test phone number."
    )
    add_test(
        doc,
        "5.1",
        "Welcome screens (customer app)",
        [
            "Open the customer app on a phone that has never opened it before (or after uninstall).",
            "Swipe through the welcome slides and tap Get started on the last slide.",
            "On the permissions screen, allow or skip location.",
        ],
        [
            "You reach the login screen.",
            "Text and images are readable; nothing is hidden under the phone status bar.",
        ],
    )
    add_test(
        doc,
        "5.2",
        "New customer profile (after first login)",
        [
            "Sign in with a fresh test number provided by your project contact (not in the table above).",
            "Fill in each step: About you → Address → Location → Solar details → Terms.",
            "Turn on location when asked; enter solar size and roof details.",
            "Tick all consent boxes and submit.",
        ],
        [
            "You land on the Home screen with bottom tabs.",
            "You can open Profile and see your details saved.",
        ],
    )

    doc.add_heading("6. Customer app tests (use Raju · 9000000401)", level=1)
    add_test(
        doc,
        "6.1",
        "Login",
        [f"Sign in with 9000000401 and OTP {TEST_OTP}."],
        ["Home screen opens without the app closing.", "Profile shows phone 9000000401."],
    )
    add_test(
        doc,
        "6.2",
        "Book a visit",
        [
            "From Home, start a new booking.",
            "Pick a date and time slot.",
            "Review the price and confirm payment (test mode—no real money).",
        ],
        ["Booking appears under the Bookings tab.", "You can open the booking detail screen."],
    )
    add_test(
        doc,
        "6.3",
        "Track a visit",
        [
            "Ask a colleague to assign a technician (see section 10) or use an existing active booking.",
            "Open the booking and check status updates.",
        ],
        ["Status changes when the technician is on the way and when the job is done.", "Map or tracking view works if shown."],
    )
    add_test(
        doc,
        "6.4",
        "Other tabs",
        [
            "Open AMC, Activity, and Profile tabs.",
            "From Profile, open Help & Support.",
        ],
        ["Each tab loads without errors.", "Help screen opens."],
    )
    add_test(
        doc,
        "6.5",
        "Wrong app",
        ["Try logging into the customer app with vendor phone 900000000201."],
        ["A message says this account is not for the customer app."],
    )

    doc.add_heading("7. Technician app tests (use Amit · 900000000301)", level=1)
    add_test(
        doc,
        "7.1",
        "Welcome screens (technician app)",
        [
            "Open the technician app for the first time on a clean install.",
            "Read the slides about field jobs and tap through to login.",
        ],
        ["Slides explain this is for technicians, not customers.", "Login screen appears."],
    )
    add_test(
        doc,
        "7.2",
        "Login",
        [f"Sign in with 900000000301 and OTP {TEST_OTP}."],
        ["Home tab opens.", "App does not crash right after login."],
    )
    add_test(
        doc,
        "7.3",
        "Do a job",
        [
            "Open the Jobs tab and select an assigned visit (coordinate with admin/vendor if empty).",
            "Tap En route when heading to the site.",
            "At the site, enter the Job Start Code shown on the customer’s booking.",
            "Complete safety checks and before/after photos.",
            "Enter the completion code and finish the visit.",
        ],
        ["Job moves to completed.", "Customer sees the visit as done."],
    )
    add_test(
        doc,
        "7.4",
        "New technician setup (optional)",
        [
            "Vendor invites a new phone number from the Team section on the vendor portal.",
            "Install technician app and sign in with that new number.",
            "Complete all onboarding steps and submit.",
            "Vendor approves the profile in the Team section.",
        ],
        ["Technician sees a “waiting for approval” screen, then full access after approval."],
    )

    doc.add_heading("8. Admin portal tests (use Priya · 900000000101)", level=1)
    doc.add_paragraph(f"Open {ADMIN_PORTAL_URL} in your browser.")
    add_test(
        doc,
        "8.1",
        "Login and dashboard",
        [f"Sign in with 900000000101 and OTP {TEST_OTP}."],
        ["Operations dashboard loads.", "No blank white screen."],
    )
    add_test(
        doc,
        "8.2",
        "Route a booking",
        [
            "Find a new customer booking in Operations or Booking routing.",
            "Assign it to vendor Gamusa Green Energy.",
        ],
        ["Assignment saves.", "Vendor can see the booking."],
    )
    add_test(
        doc,
        "8.3",
        "Vendors and technicians",
        [
            "Open vendor list and technician directory.",
            "Confirm Gamusa and Bharat Sun show as approved.",
        ],
        ["Lists load and show expected partners."],
    )

    doc.add_heading("9. Vendor portal tests (use Gamusa · 900000000201)", level=1)
    doc.add_paragraph(f"Open {VENDOR_PORTAL_URL} in your browser.")
    add_test(
        doc,
        "9.1",
        "Login and overview",
        [f"Sign in with 900000000201 and OTP {TEST_OTP}."],
        ["Overview page shows summary numbers.", "Side or top menu is visible."],
    )
    add_test(
        doc,
        "9.2",
        "Check all main sections",
        [
            "Click through: Overview, Operations, Insights, Finance, Team, Bookings, Service coverage.",
        ],
        ["Each section opens without error."],
    )
    add_test(
        doc,
        "9.3",
        "Assign a technician",
        [
            "In Operations or Bookings, open a visit for customer Raju (401).",
            "Assign technician Amit Das (301).",
        ],
        ["Amit sees the job in the technician app Jobs tab."],
    )

    doc.add_heading("10. Support desk tests (use Ananya · 900000000111)", level=1)
    doc.add_paragraph(f"Open {SUPPORT_PORTAL_URL} in your browser.")
    add_test(
        doc,
        "10.1",
        "Login",
        [f"Sign in with 900000000111 and OTP {TEST_OTP}."],
        ["Insights page loads."],
    )
    add_test(
        doc,
        "10.2",
        "Find a customer",
        [
            "Open Inbox or Search.",
            "Search for phone 9000000401 or a booking ID from a test visit.",
        ],
        ["Customer or booking details appear."],
    )

    doc.add_heading("11. Full journey (do this once per release)", level=1)
    doc.add_paragraph("One person plays customer, others play admin, vendor, and technician:")
    journey = [
        "Customer Raju (401) books a cleaning visit in the customer app.",
        f"Admin Priya (101) logs in at {ADMIN_PORTAL_URL} and routes the booking to Gamusa.",
        f"Vendor Gamusa (201) logs in at {VENDOR_PORTAL_URL} and assigns Amit (301).",
        "Technician Amit (301) completes the visit in the technician app.",
        "Customer Raju checks Bookings and Activity—the visit shows as completed.",
        f"Support Ananya (111) can find the booking at {SUPPORT_PORTAL_URL} if needed.",
    ]
    for step in journey:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("12. Sign-off checklist", level=1)
    doc.add_paragraph("Tick each item when satisfied:")
    for item in [
        "Customer APK installs and opens.",
        "Technician APK installs and opens.",
        "All three web portals open in the browser.",
        f"OTP {TEST_OTP} works on every app.",
        "Customer can book a visit.",
        "Admin can route a booking.",
        "Vendor can assign a technician.",
        "Technician can complete a visit.",
        "Customer sees the completed visit.",
        "No crash right after login on Android.",
        "New customer welcome + profile flow tested once (section 5).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)

    doc.add_heading("13. When something goes wrong", level=1)
    for item in [
        "Note the app name, test phone used, and what you tapped just before the problem.",
        "Take a screenshot or screen recording.",
        "If login fails for everyone, tell your project contact—the test environment may need a refresh.",
        "To see first-time welcome slides again, uninstall the app and reinstall the APK.",
        "To clear old bookings and try again, ask your project contact to reset test data.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    note = doc.add_paragraph(
        f"Portal URLs: Admin {ADMIN_PORTAL_URL} · Vendor {VENDOR_PORTAL_URL} · Support {SUPPORT_PORTAL_URL}"
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
